import io
import os
from pypdf import PdfReader
import docx

def extract_text_from_pdf(file_input) -> str:
    """Extract raw text from a PDF file path or file-like object stream."""
    try:
        reader = PdfReader(file_input)
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF document: {e}")

def extract_text_from_docx(file_input) -> str:
    """Extract raw text from a DOCX file path or file-like object stream."""
    try:
        doc = docx.Document(file_input)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Also extract table text to avoid missing data
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
                    
        return "\n".join(text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX document: {e}")

def extract_text_from_txt(file_input) -> str:
    """Extract raw text from a TXT file path or file-like stream."""
    try:
        if isinstance(file_input, str):
            with open(file_input, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            # Handle stream
            content = file_input.read()
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="ignore")
            return content
    except Exception as e:
        raise ValueError(f"Failed to read TXT file: {e}")

def extract_text(file_path_or_stream, file_extension: str = None) -> str:
    """General helper to route text extraction based on file extension."""
    if file_extension is None and isinstance(file_path_or_stream, str):
        _, ext = os.path.splitext(file_path_or_stream)
        file_extension = ext.lower().lstrip(".")
        
    if not file_extension:
        raise ValueError("File extension could not be resolved.")
        
    ext = file_extension.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_from_pdf(file_path_or_stream)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_path_or_stream)
    elif ext in ("txt", "md"):
        return extract_text_from_txt(file_path_or_stream)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
